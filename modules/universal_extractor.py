"""
Universal Text Extractor for Consciousness Recognition System

Supports multiple e-book formats: PDF, EPUB, AZW3, MOBI, TXT, DOCX, and more.
Provides unified interface for text extraction from various spiritual text sources.
"""

import os
import tempfile
import zipfile
import xml.etree.ElementTree as ET
from typing import Dict, Any, Optional, List, Tuple
from pathlib import Path
import re
import subprocess


class UniversalTextExtractor:
    """Universal text extractor supporting multiple e-book formats."""
    
    def __init__(self):
        """Initialize the universal extractor."""
        self.supported_formats = {
            '.pdf': self._extract_pdf,
            '.epub': self._extract_epub,
            '.azw3': self._extract_azw3,
            '.mobi': self._extract_mobi,
            '.txt': self._extract_txt,
            '.docx': self._extract_docx,
            '.doc': self._extract_doc,
            '.rtf': self._extract_rtf,
            '.html': self._extract_html,
            '.htm': self._extract_html
        }
        
        # Check available dependencies
        self.available_extractors = self._check_dependencies()
    
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from any supported format.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dict with extraction results
        """
        try:
            file_ext = Path(file_path).suffix.lower()
            
            if file_ext not in self.supported_formats:
                return {
                    'success': False,
                    'text': '',
                    'format': file_ext,
                    'error': f'Unsupported format: {file_ext}',
                    'metadata': {}
                }
            
            # Check if extractor is available
            if file_ext not in self.available_extractors:
                return {
                    'success': False,
                    'text': '',
                    'format': file_ext,
                    'error': f'Extractor not available for {file_ext}. Missing dependencies.',
                    'metadata': {}
                }
            
            # Extract using appropriate method
            extractor_func = self.supported_formats[file_ext]
            result = extractor_func(file_path)
            
            # Add format info
            result['format'] = file_ext
            
            return result
            
        except Exception as e:
            return {
                'success': False,
                'text': '',
                'format': Path(file_path).suffix.lower(),
                'error': f'Extraction failed: {str(e)}',
                'metadata': {}
            }
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats."""
        return list(self.available_extractors.keys())
    
    def get_format_info(self, file_path: str) -> Dict[str, Any]:
        """Get information about file format and extraction capabilities."""
        file_ext = Path(file_path).suffix.lower()
        
        format_info = {
            '.pdf': {
                'name': 'Portable Document Format',
                'description': 'Standard document format, may contain text or images',
                'extraction_quality': 'High for text-based, Medium for scanned',
                'metadata_support': True
            },
            '.epub': {
                'name': 'Electronic Publication',
                'description': 'Standard e-book format with structured content',
                'extraction_quality': 'Excellent',
                'metadata_support': True
            },
            '.azw3': {
                'name': 'Amazon Kindle Format',
                'description': 'Amazon proprietary e-book format',
                'extraction_quality': 'Good (requires calibre)',
                'metadata_support': True
            },
            '.mobi': {
                'name': 'Mobipocket',
                'description': 'Legacy e-book format, predecessor to AZW',
                'extraction_quality': 'Good (requires calibre)',
                'metadata_support': True
            },
            '.txt': {
                'name': 'Plain Text',
                'description': 'Simple text file',
                'extraction_quality': 'Perfect',
                'metadata_support': False
            },
            '.docx': {
                'name': 'Microsoft Word Document',
                'description': 'Modern Word document format',
                'extraction_quality': 'Excellent',
                'metadata_support': True
            }
        }
        
        return {
            'format': file_ext,
            'supported': file_ext in self.available_extractors,
            'info': format_info.get(file_ext, {'name': 'Unknown', 'description': 'Unknown format'})
        }
    
    def _check_dependencies(self) -> Dict[str, str]:
        """Check which extractors are available based on installed dependencies."""
        available = {}
        
        # PDF extraction
        try:
            import fitz
            available['.pdf'] = 'PyMuPDF'
        except ImportError:
            try:
                import pdfplumber
                available['.pdf'] = 'pdfplumber'
            except ImportError:
                pass
        
        # EPUB extraction
        try:
            import ebooklib
            available['.epub'] = 'ebooklib'
        except ImportError:
            pass
        
        # Text files (always available)
        available['.txt'] = 'built-in'
        
        # DOCX extraction
        try:
            import docx
            available['.docx'] = 'python-docx'
        except ImportError:
            pass
        
        # HTML extraction
        try:
            import bs4
            available['.html'] = 'BeautifulSoup'
            available['.htm'] = 'BeautifulSoup'
        except ImportError:
            pass
        
        # Check for calibre (for AZW3/MOBI)
        try:
            result = subprocess.run(['ebook-convert', '--version'], 
                                  capture_output=True, text=True, timeout=5)
            if result.returncode == 0:
                available['.azw3'] = 'calibre'
                available['.mobi'] = 'calibre'
        except (subprocess.TimeoutExpired, FileNotFoundError):
            pass
        
        return available
    
    def _extract_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF files."""
        try:
            # Try PyMuPDF first
            try:
                import fitz
                doc = fitz.open(file_path)
                text = ""
                metadata = {
                    'page_count': len(doc),
                    'title': doc.metadata.get('title', ''),
                    'author': doc.metadata.get('author', ''),
                    'subject': doc.metadata.get('subject', '')
                }
                
                for page in doc:
                    text += page.get_text() + "\n"
                
                doc.close()
                
                return {
                    'success': True,
                    'text': text,
                    'metadata': metadata,
                    'extractor': 'PyMuPDF'
                }
                
            except ImportError:
                # Fallback to pdfplumber
                import pdfplumber
                
                text = ""
                with pdfplumber.open(file_path) as pdf:
                    metadata = {
                        'page_count': len(pdf.pages),
                        'title': pdf.metadata.get('Title', ''),
                        'author': pdf.metadata.get('Author', ''),
                        'subject': pdf.metadata.get('Subject', '')
                    }
                    
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                
                return {
                    'success': True,
                    'text': text,
                    'metadata': metadata,
                    'extractor': 'pdfplumber'
                }
                
        except Exception as e:
            return {
                'success': False,
                'text': '',
                'error': f'PDF extraction failed: {str(e)}',
                'metadata': {}
            }
    
    def _extract_epub(self, file_path: str) -> Dict[str, Any]:
        """Extract text from EPUB files."""
        try:
            import ebooklib
            from ebooklib import epub
            from bs4 import BeautifulSoup
            
            book = epub.read_epub(file_path)
            
            # Extract metadata
            metadata = {
                'title': book.get_metadata('DC', 'title')[0][0] if book.get_metadata('DC', 'title') else '',
                'author': book.get_metadata('DC', 'creator')[0][0] if book.get_metadata('DC', 'creator') else '',
                'language': book.get_metadata('DC', 'language')[0][0] if book.get_metadata('DC', 'language') else '',
                'publisher': book.get_metadata('DC', 'publisher')[0][0] if book.get_metadata('DC', 'publisher') else ''
            }
            
            # Extract text from all chapters
            text = ""
            chapter_count = 0
            
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    chapter_count += 1
                    soup = BeautifulSoup(item.get_content(), 'html.parser')
                    
                    # Remove script and style elements
                    for script in soup(["script", "style"]):
                        script.decompose()
                    
                    # Get text and clean it
                    chapter_text = soup.get_text()
                    lines = (line.strip() for line in chapter_text.splitlines())
                    chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                    chapter_text = ' '.join(chunk for chunk in chunks if chunk)
                    
                    text += chapter_text + "\n\n"
            
            metadata['chapter_count'] = chapter_count
            
            return {
                'success': True,
                'text': text,
                'metadata': metadata,
                'extractor': 'ebooklib'
            }
            
        except Exception as e:
            return {
                'success': False,
                'text': '',
                'error': f'EPUB extraction failed: {str(e)}',
                'metadata': {}
            }
    
    def _extract_azw3(self, file_path: str) -> Dict[str, Any]:
        """Extract text from AZW3 files using calibre."""
        return self._extract_with_calibre(file_path, 'AZW3')
    
    def _extract_mobi(self, file_path: str) -> Dict[str, Any]:
        """Extract text from MOBI files using calibre."""
        return self._extract_with_calibre(file_path, 'MOBI')
    
    def _extract_with_calibre(self, file_path: str, format_name: str) -> Dict[str, Any]:
        """Extract text using calibre's ebook-convert."""
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                output_path = os.path.join(temp_dir, 'output.txt')
                
                # Convert to text using calibre
                cmd = ['ebook-convert', file_path, output_path]
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
                
                if result.returncode != 0:
                    return {
                        'success': False,
                        'text': '',
                        'error': f'{format_name} conversion failed: {result.stderr}',
                        'metadata': {}
                    }
                
                # Read extracted text
                with open(output_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                
                # Extract basic metadata from calibre output
                metadata = {
                    'extractor': 'calibre',
                    'original_format': format_name
                }
                
                # Try to extract metadata from calibre
                meta_cmd = ['ebook-meta', file_path]
                meta_result = subprocess.run(meta_cmd, capture_output=True, text=True, timeout=30)
                
                if meta_result.returncode == 0:
                    # Parse metadata from calibre output
                    meta_lines = meta_result.stdout.split('\n')
                    for line in meta_lines:
                        if ':' in line:
                            key, value = line.split(':', 1)
                            key = key.strip().lower()
                            value = value.strip()
                            
                            if key in ['title', 'author', 'publisher', 'language']:
                                metadata[key] = value
                
                return {
                    'success': True,
                    'text': text,
                    'metadata': metadata,
                    'extractor': 'calibre'
                }
                
        except subprocess.TimeoutExpired:
            return {
                'success': False,
                'text': '',
                'error': f'{format_name} extraction timed out',
                'metadata': {}
            }
        except Exception as e:
            return {
                'success': False,
                'text': '',
                'error': f'{format_name} extraction failed: {str(e)}',
                'metadata': {}
            }
    
    def _extract_txt(self, file_path: str) -> Dict[str, Any]:
        """Extract text from plain text files."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    
                    # Get file stats
                    stat = os.stat(file_path)
                    metadata = {
                        'encoding': encoding,
                        'file_size': stat.st_size,
                        'line_count': text.count('\n') + 1
                    }
                    
                    return {
                        'success': True,
                        'text': text,
                        'metadata': metadata,
                        'extractor': 'built-in'
                    }
                    
                except UnicodeDecodeError:
                    continue
            
            return {
                'success': False,
                'text': '',
                'error': 'Could not decode text file with any supported encoding',
                'metadata': {}
            }
            
        except Exception as e:
            return {
                'success': False,
                'text': '',
                'error': f'Text extraction failed: {str(e)}',
                'metadata': {}
            }
    
    def _extract_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX files."""
        try:
            import docx
            
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract metadata
            metadata = {
                'title': doc.core_properties.title or '',
                'author': doc.core_properties.author or '',
                'subject': doc.core_properties.subject or '',
                'paragraph_count': len(doc.paragraphs)
            }
            
            return {
                'success': True,
                'text': text,
                'metadata': metadata,
                'extractor': 'python-docx'
            }
            
        except Exception as e:
            return {
                'success': False,
                'text': '',
                'error': f'DOCX extraction failed: {str(e)}',
                'metadata': {}
            }
    
    def _extract_doc(self, file_path: str) -> Dict[str, Any]:
        """Extract text from legacy DOC files."""
        try:
            # Try using python-docx2txt if available
            try:
                import docx2txt
                text = docx2txt.process(file_path)
                
                return {
                    'success': True,
                    'text': text,
                    'metadata': {'extractor': 'docx2txt'},
                    'extractor': 'docx2txt'
                }
                
            except ImportError:
                # Fallback: suggest conversion
                return {
                    'success': False,
                    'text': '',
                    'error': 'DOC format requires docx2txt library or manual conversion to DOCX',
                    'metadata': {}
                }
                
        except Exception as e:
            return {
                'success': False,
                'text': '',
                'error': f'DOC extraction failed: {str(e)}',
                'metadata': {}
            }
    
    def _extract_rtf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from RTF files."""
        try:
            # Try using striprtf if available
            try:
                from striprtf.striprtf import rtf_to_text
                
                with open(file_path, 'r', encoding='utf-8') as f:
                    rtf_content = f.read()
                
                text = rtf_to_text(rtf_content)
                
                return {
                    'success': True,
                    'text': text,
                    'metadata': {'extractor': 'striprtf'},
                    'extractor': 'striprtf'
                }
                
            except ImportError:
                return {
                    'success': False,
                    'text': '',
                    'error': 'RTF format requires striprtf library',
                    'metadata': {}
                }
                
        except Exception as e:
            return {
                'success': False,
                'text': '',
                'error': f'RTF extraction failed: {str(e)}',
                'metadata': {}
            }
    
    def _extract_html(self, file_path: str) -> Dict[str, Any]:
        """Extract text from HTML files."""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text
            text = soup.get_text()
            
            # Clean up text
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            # Extract basic metadata
            metadata = {
                'title': soup.title.string if soup.title else '',
                'extractor': 'BeautifulSoup'
            }
            
            return {
                'success': True,
                'text': text,
                'metadata': metadata,
                'extractor': 'BeautifulSoup'
            }
            
        except Exception as e:
            return {
                'success': False,
                'text': '',
                'error': f'HTML extraction failed: {str(e)}',
                'metadata': {}
            }

